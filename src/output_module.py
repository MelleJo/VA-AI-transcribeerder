# src/output_module.py

import streamlit as st
from src import utils
from st_copy_to_clipboard import st_copy_to_clipboard
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from datetime import datetime
import markdown2
from io import BytesIO
from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from xhtml2pdf import pisa
import logging
import traceback

logging.basicConfig(level=logging.INFO)

def markdown_to_html(markdown_text):
    return markdown2.markdown(markdown_text)

def create_pdf(markdown_content):
    html_content = markdown2.markdown(markdown_content)
    
    css = """
    <style>
    body {
        font-family: Arial, sans-serif;
        font-size: 11pt;
        line-height: 1.6;
    }
    p {
        margin-bottom: 10px;
    }
    strong {
        font-weight: bold;
    }
    </style>
    """
    
    html = f"""
    <html>
    <head>
    {css}
    </head>
    <body>
    {html_content}
    </body>
    </html>
    """
    
    pdf_buffer = BytesIO()
    try:
        pisa_status = pisa.CreatePDF(html, dest=pdf_buffer)
        if pisa_status.err:
            logging.error(f"PDF creation error: {pisa_status.err}")
            return None
        pdf_buffer.seek(0)
        return pdf_buffer
    except Exception as e:
        logging.error(f"Error creating PDF: {str(e)}")
        logging.error(traceback.format_exc())
        return None

def create_docx(markdown_content):
    try:
        doc = Document()
        
        # Add styles
        styles = doc.styles
        style = styles.add_style('Normal', WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = 'Arial'
        style.font.size = Pt(11)
        
        # Split the markdown content into paragraphs
        paragraphs = markdown_content.split('\n\n')
        
        for para in paragraphs:
            if para.startswith('#'):
                # Handle headers
                level = para.count('#')
                text = para.strip('#').strip()
                doc.add_heading(text, level=min(level, 9))
            elif para.startswith('- '):
                # Handle unordered lists
                doc.add_paragraph(para.strip('- '), style='List Bullet')
            elif para.startswith('1. '):
                # Handle ordered lists
                doc.add_paragraph(para[3:], style='List Number')
            else:
                # Regular paragraph
                doc.add_paragraph(para, style='Normal')
        
        docx_buffer = BytesIO()
        doc.save(docx_buffer)
        docx_buffer.seek(0)
        return docx_buffer
    except Exception as e:
        logging.error(f"Error creating DOCX: {str(e)}")
        logging.error(traceback.format_exc())
        return None

def send_feedback_email(transcript, summary, feedback, additional_feedback, user_name):
    sender_email = st.secrets["email"]["username"]
    receiver_email = st.secrets["email"]["receiving_email"]
    password = st.secrets["email"]["password"]
    smtp_server = st.secrets["email"]["smtp_server"]
    smtp_port = st.secrets["email"]["smtp_port"]

    message = MIMEMultipart("alternative")
    message["Subject"] = f"Gesprekssamenvatter Feedback - {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
    message["From"] = sender_email
    message["To"] = receiver_email

    text = f"""
    Naam: {user_name}
    Feedback: {feedback}
    Aanvullende feedback: {additional_feedback}

    Transcript:
    {transcript}

    Samenvatting:
    {summary}
    """

    html = f"""
    <html>
    <body>
        <h2>Gesprekssamenvatter Feedback</h2>
        <p><strong>Naam:</strong> {user_name}</p>
        <p><strong>Feedback:</strong> {feedback}</p>
        <p><strong>Aanvullende feedback:</strong> {additional_feedback}</p>
        
        <h3>Transcript:</h3>
        <pre>{transcript}</pre>
        
        <h3>Samenvatting:</h3>
        <pre>{summary}</pre>
    </body>
    </html>
    """

    part1 = MIMEText(text, "plain")
    part2 = MIMEText(html, "html")

    message.attach(part1)
    message.attach(part2)

    try:
        with smtplib.SMTP(smtp_server, smtp_port) as server:
            server.starttls()
            server.login(sender_email, password)
            server.sendmail(sender_email, receiver_email, message.as_string())
        return True
    except Exception as e:
        st.error(f"Er is een fout opgetreden bij het verzenden van de e-mail: {str(e)}")
        return False

def render_output():
    st.header("Stap 3: Samenvatting")

    if not st.session_state.summary:
        st.warning("Er is nog geen samenvatting gegenereerd. Voltooi eerst de vorige stappen.")
        return

    st.markdown("### Gegenereerde samenvatting")
    st.markdown(st.session_state.summary)

    col1, col2, col3 = st.columns(3)

    with col1:
        if st_copy_to_clipboard(st.session_state.summary):
            st.success("Gekopieerd!")

    with col2:
        # PDF download button
        pdf_buffer = create_pdf(st.session_state.summary)
        if pdf_buffer:
            st.download_button(
                label="Download als PDF",
                data=pdf_buffer,
                file_name="samenvatting.pdf",
                mime="application/pdf"
            )
        else:
            st.error("PDF genereren mislukt.")

    with col3:
        # DOCX download button
        docx_buffer = create_docx(st.session_state.summary)
        if docx_buffer:
            st.download_button(
                label="Download als DOCX",
                data=docx_buffer,
                file_name="samenvatting.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        else:
            st.error("DOCX genereren mislukt.")

    st.markdown("### Feedback")
    with st.form(key="feedback_form"):
        user_name = st.text_input("Uw naam (verplicht bij feedback):")
        feedback = st.radio("Was deze samenvatting nuttig?", ["Positief", "Negatief"])
        additional_feedback = st.text_area("Laat aanvullende feedback achter:")
        submit_button = st.form_submit_button(label="Verzend feedback")

        if submit_button:
            if not user_name:
                st.warning("Naam is verplicht bij het geven van feedback.", icon="⚠️")
            else:
                success = send_feedback_email(
                    transcript=st.session_state.input_text,
                    summary=st.session_state.summary,
                    feedback=feedback,
                    additional_feedback=additional_feedback,
                    user_name=user_name
                )
                if success:
                    st.success("Bedankt voor de feedback!")
                else:
                    st.error("Er is een fout opgetreden bij het verzenden van de feedback. Probeer het later opnieuw.")

    if st.button("Start nieuwe samenvatting"):
        for key in ['input_text', 'selected_prompt', 'summary']:
            if key in st.session_state:
                del st.session_state[key]
        st.session_state.step = 1
        st.rerun()