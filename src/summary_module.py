import streamlit as st
from openai import OpenAI
from src.config import SUMMARY_MODEL, MAX_TOKENS, TEMPERATURE, TOP_P, FREQUENCY_PENALTY, PRESENCE_PENALTY
from src.history_module import add_to_history
from src.utils import load_prompts, get_prompt_content
from st_copy_to_clipboard import st_copy_to_clipboard
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from datetime import datetime
import base64
import io
from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.pagesizes import letter
import markdown2
import re

client = OpenAI(api_key=st.secrets["OPENAI_API_KEY"])

def markdown_to_html(markdown_text):
    return markdown2.markdown(markdown_text)

def generate_summary(input_text, base_prompt, selected_prompt):
    try:
        full_prompt = f"{base_prompt}\n\n{selected_prompt}"
        
        response = client.chat.completions.create(
            model=SUMMARY_MODEL,
            messages=[
                {"role": "system", "content": full_prompt},
                {"role": "user", "content": input_text}
            ],
            max_tokens=MAX_TOKENS,
            temperature=TEMPERATURE,
            top_p=TOP_P,
            frequency_penalty=FREQUENCY_PENALTY,
            presence_penalty=PRESENCE_PENALTY,
            n=1,
            stop=None
        )
        summary = response.choices[0].message.content.strip()
        if not summary:
            raise ValueError("Generated summary is empty")
        return summary
    except Exception as e:
        st.error(f"Er is een fout opgetreden bij het genereren van de samenvatting: {str(e)}")
        return None

def customize_summary(current_summary, customization_request, transcript):
    try:
        response = client.chat.completions.create(
            model=SUMMARY_MODEL,
            messages=[
                {"role": "system", "content": "Je bent een AI-assistent die samenvattingen aanpast op basis van specifieke verzoeken. Behoud de essentie van de oorspronkelijke samenvatting, maar pas deze aan volgens het verzoek van de gebruiker."},
                {"role": "user", "content": f"Oorspronkelijke samenvatting:\n\n{current_summary}\n\nTranscript:\n\n{transcript}\n\nAanpassingsverzoek: {customization_request}\n\nPas de samenvatting aan volgens dit verzoek."}
            ],
            max_tokens=MAX_TOKENS,
            temperature=TEMPERATURE,
            top_p=TOP_P,
            frequency_penalty=FREQUENCY_PENALTY,
            presence_penalty=PRESENCE_PENALTY,
            n=1,
            stop=None
        )
        customized_summary = response.choices[0].message.content.strip()
        if not customized_summary:
            raise ValueError("Aangepaste samenvatting is leeg")
        return customized_summary
    except Exception as e:
        st.error(f"Er is een fout opgetreden bij het aanpassen van de samenvatting: {str(e)}")
        return None

def render_summary_buttons(summary, button_key_prefix):
    # Remove HTML tags for clipboard copy
    clean_summary = re.sub('<[^<]+?>', '', summary)
    
    # Using st_copy_to_clipboard for the clipboard functionality
    if st_copy_to_clipboard(clean_summary, key=f"{button_key_prefix}_copy"):
        st.success(f"Samenvatting gekopieerd naar klembord!")

    col1, col2 = st.columns(2)
    with col1:
        b64_docx = export_to_docx(summary)
        href_docx = f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64_docx}" download="samenvatting_{button_key_prefix}.docx">Download als Word-document</a>'
        st.markdown(href_docx, unsafe_allow_html=True)

    with col2:
        b64_pdf = export_to_pdf(summary)
        href_pdf = f'<a href="data:application/pdf;base64,{b64_pdf}" download="samenvatting_{button_key_prefix}.pdf">Download als PDF</a>'
        st.markdown(href_pdf, unsafe_allow_html=True)

def render_summary_generation():
    st.header("Stap 4: Samenvatting")
    
    prompts = load_prompts()
    base_prompt = prompts.get('base_prompt.txt', '')
    prompt_name = st.session_state.get('selected_prompt', 'None')
    selected_prompt = get_prompt_content(prompt_name)
    
    input_text_length = len(st.session_state.get('input_text', ''))

    # Create debug info using Streamlit's native expander
    with st.expander("Debug Info"):
        st.write(f"Geselecteerde prompt: {prompt_name}")
        st.write(f"Basis prompt geladen: {'Ja' if base_prompt else 'Nee'}")
        st.write(f"Geselecteerde prompt geladen: {'Ja' if selected_prompt else 'Nee'}")
        st.write(f"Lengte invoertekst: {input_text_length} tekens")
        st.write(f"Max tokens voor samenvatting: {MAX_TOKENS}")

    if not st.session_state.input_text:
        st.warning("Er is geen tekst om samen te vatten. Ga terug naar de vorige stappen om tekst in te voeren.")
        return

    summary_placeholder = st.empty()

    if not st.session_state.summary:
        with st.spinner("Samenvatting wordt gegenereerd..."):
            summary = generate_summary(st.session_state.input_text, base_prompt, selected_prompt)
            if summary:
                st.session_state.summary = summary
                add_to_history(prompt_name, st.session_state.input_text, summary)
                summary_placeholder.success("Samenvatting succesvol gegenereerd! Ik hoor graag feedback (negatief én positief!) via de feedbacktool onderin het scherm.")
            else:
                summary_placeholder.error("Samenvatting genereren mislukt. Probeer het opnieuw.")

    if st.session_state.summary:
        st.markdown("### Gegenereerde Samenvatting")
        st.markdown(st.session_state.summary)

        # Custom CSS to style the st_copy_to_clipboard button
        st.markdown("""
            <style>
            .stButton button {
                background-color: #4CAF50;
                color: white;
                padding: 15px 32px;
                font-size: 16px;
                border-radius: 8px;
                border: none;
                cursor: pointer;
            }
            </style>
        """, unsafe_allow_html=True)

        render_summary_buttons(st.session_state.summary, "initial")

        # Add a button to expand the customization section
        if st.button("Pas samenvatting aan", key="customize_summary_button"):
            st.session_state.show_customization = True

        # Show the customization section if the button was clicked
        if st.session_state.get('show_customization', False):
            st.markdown("### Pas de samenvatting aan")
            customization_request = st.text_area("Voer hier uw aanpassingsverzoek in (bijv. 'Maak het korter', 'Voeg meer details toe over X', 'Maak het formeler'):", key="customization_request")
            if st.button("Pas samenvatting aan", key="apply_customization_button"):
                with st.spinner("Samenvatting wordt aangepast..."):
                    customized_summary = customize_summary(st.session_state.summary, customization_request, st.session_state.input_text)
                    if customized_summary:
                        st.session_state.revised_summary = customized_summary
                        st.success("Samenvatting succesvol aangepast!")
                        st.markdown("### Aangepaste Samenvatting")
                        st.markdown(customized_summary)
                        render_summary_buttons(customized_summary, "revised")
                    else:
                        st.error("Aanpassing van de samenvatting mislukt. Probeer het opnieuw.")

        if st.button("Genereer Nieuwe Samenvatting", key="generate_new_summary_button"):
            st.session_state.summary = None
            st.session_state.revised_summary = None
            st.session_state.show_customization = False
            st.rerun()

        st.markdown("### Feedback")
        with st.form(key="feedback_form"):
            user_name = st.text_input("Uw naam (verplicht bij feedback):", key="feedback_name")
            feedback = st.radio("Was deze samenvatting nuttig?", ["Positief", "Negatief"], key="feedback_rating")
            additional_feedback = st.text_area("Laat aanvullende feedback achter:", key="additional_feedback")
            submit_button = st.form_submit_button(label="Verzend feedback")

            if submit_button:
                if not user_name:
                    st.warning("Naam is verplicht bij het geven van feedback.", icon="⚠️")
                else:
                    success = send_feedback_email(
                        transcript=st.session_state.input_text,
                        summary=st.session_state.summary,
                        revised_summary=st.session_state.get('revised_summary', 'Geen aangepaste samenvatting'),
                        feedback=feedback,
                        additional_feedback=additional_feedback,
                        user_name=user_name,
                        selected_prompt=prompt_name
                    )
                    if success:
                        st.success("Bedankt voor uw feedback!")
                    else:
                        st.error("Er is een fout opgetreden bij het verzenden van de feedback. Probeer het later opnieuw.")

        if st.button("Start Nieuwe Samenvatting", key="start_new_summary_button"):
            for key in ['input_text', 'selected_prompt', 'summary', 'revised_summary', 'show_customization']:
                if key in st.session_state:
                    del st.session_state[key]
            st.session_state.step = 1
            st.rerun()

def send_feedback_email(transcript, summary, revised_summary, feedback, additional_feedback, user_name, selected_prompt):
    try:
        sender_email = st.secrets["email"]["username"]
        receiver_email = st.secrets["email"]["receiving_email"]
        password = st.secrets["email"]["password"]
        smtp_server = st.secrets["email"]["smtp_server"]
        smtp_port = st.secrets["email"]["smtp_port"]

        message = MIMEMultipart("alternative")
        message["Subject"] = f"Gesprekssamenvatter Feedback - {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
        message["From"] = sender_email
        message["To"] = receiver_email

        text = f"""
        Naam: {user_name}
        Feedback: {feedback}
        Aanvullende feedback: {additional_feedback}
        Geselecteerde prompt: {selected_prompt}

        Transcript:
        {transcript}

        Oorspronkelijke Samenvatting:
        {summary}

        Aangepaste Samenvatting:
        {revised_summary}
        """

        html = f"""
        <html>
        <body>
            <h2>Gesprekssamenvatter Feedback</h2>
            <p><strong>Naam:</strong> {user_name}</p>
            <p><strong>Feedback:</strong> {feedback}</p>
            <p><strong>Aanvullende feedback:</strong> {additional_feedback}</p>
            <p><strong>Geselecteerde prompt:</strong> {selected_prompt}</p>
            
            <h3>Transcript:</h3>
            <pre>{transcript}</pre>
            
            <h3>Oorspronkelijke Samenvatting:</h3>
            <pre>{summary}</pre>

            <h3>Aangepaste Samenvatting:</h3>
            <pre>{revised_summary}</pre>
        </body>
        </html>
        """

        part1 = MIMEText(text, "plain")
        part2 = MIMEText(html, "html")

        message.attach(part1)
        message.attach(part2)

        with smtplib.SMTP(smtp_server, smtp_port) as server:
            server.starttls()
            server.login(sender_email, password)
            server.sendmail(sender_email, receiver_email, message.as_string())
        return True
    except Exception as e:
        st.error(f"Er is een fout opgetreden bij het verzenden van de e-mail: {str(e)}")
        return False

def export_to_docx(summary):
    doc = Document()
    styles = doc.styles
    try:
        style = styles.add_style('Body Text', WD_STYLE_TYPE.PARAGRAPH)
    except ValueError:
        # If the style already exists, just get it
        style = styles['Body Text']
    style.font.size = Pt(11)
    
    paragraphs = markdown_to_html(summary).split('<p>')
    for p in paragraphs:
        if p.strip():
            para = doc.add_paragraph()
            para.style = 'Body Text'
            run = para.add_run(p.replace('</p>', '').strip())
            if '<strong>' in p:
                run.bold = True
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    b64 = base64.b64encode(buffer.getvalue()).decode()
    return b64

def export_to_pdf(summary):
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    
    story = []
    paragraphs = markdown_to_html(summary).split('<p>')
    for p in paragraphs:
        if p.strip():
            if '<strong>' in p:
                p = p.replace('<strong>', '<b>').replace('</strong>', '</b>')
            story.append(Paragraph(p.replace('</p>', '').strip(), styles['BodyText']))
            story.append(Spacer(1, 12))
    
    doc.build(story)
    buffer.seek(0)
    
    b64 = base64.b64encode(buffer.getvalue()).decode()
    return b64
